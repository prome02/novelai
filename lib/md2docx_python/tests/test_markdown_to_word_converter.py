import os
import pytest
from docx import Document
# from src.markdown_to_word_converter import markdown_to_word

def test_dummy():
    pass
#
# @pytest.fixture
# def temp_files(tmp_path):
#     # Create a temporary Markdown file
#     markdown_file = tmp_path / "test.md"
#     markdown_file.write_text(
#         "# Heading 1\n\n## Heading 2\n\nThis is **bold** text and *italic* text.\n\n- Item 1\n- Item 2\n\n1. Item A\n2. Item B")
#
#     # Define a path for the Word file
#     word_file = tmp_path / "test.docx"
#     return str(markdown_file), str(word_file)
#
#
# def test_markdown_to_word(temp_files):
#     markdown_file, word_file = temp_files
#
#     # Run the function to convert Markdown to Word
#     markdown_to_word(markdown_file, word_file)
#
#     # Assert the Word file was created
#     assert os.path.exists(word_file), "Word file was not created."
#
#     # Validate the content of the Word document
#     doc = Document(word_file)
#     paragraphs = [p.text for p in doc.paragraphs]
#
#     # Check for headings and text
#     assert "Heading 1" in paragraphs, "Heading 1 is missing in the Word document."
#     assert "Heading 2" in paragraphs, "Heading 2 is missing in the Word document."
#     assert "This is bold text and italic text." in paragraphs, "Paragraph text is missing or incorrect."
#
#     # Validate list items
#     list_items = [p.text for p in doc.paragraphs if p.style.name.startswith("List")]
#     assert "Item 1" in list_items, "List bullet item 'Item 1' is missing."
#     assert "Item 2" in list_items, "List bullet item 'Item 2' is missing."
#     assert "Item A" in list_items, "List numbered item 'Item A' is missing."
#     assert "Item B" in list_items, "List numbered item 'Item B' is missing."
