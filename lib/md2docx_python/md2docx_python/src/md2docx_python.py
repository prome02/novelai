import markdown
from docx import Document
from bs4 import BeautifulSoup

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def markdown_to_word(markdown_file, word_file, word_template=None):
    # Reading the Markdown file
    with open(markdown_file, 'r', encoding='utf-8') as file:
        markdown_content = file.read().replace("\n\n---\n\n", "\n\n<hr/><br /><br />\n\n")

    # Converting Markdown to HTML
    html_content = markdown.markdown(markdown_content)

    # Creating a new Word Document
    doc = Document(word_template)

    # Converting HTML to text and add it to the Word Document
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Adding content to the Word Document
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_page_break()
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'hr':
            paragraph = doc.add_paragraph()
            insertHR(paragraph)
        elif element.name == 'p':
            paragraph = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong':
                    paragraph.add_run(child.text).bold = True
                elif child.name == 'em':
                    paragraph.add_run(child.text).italic = True
                else:
                    paragraph.add_run(child)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
    
    doc.save(word_file)
